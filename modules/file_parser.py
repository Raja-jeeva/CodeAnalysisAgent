import io
import os
import re
from typing import List, Dict, Any
from docx import Document

from .logger import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {
    ".py", ".js", ".ts", ".java", ".cs", ".cpp", ".c", ".h", ".hpp", ".go", ".rb", ".php", ".swift", ".kt", ".rs", ".m", ".mm", ".scala"
}
MAX_FILE_SIZE_BYTES = 800 * 1024  # 800KB per file safety limit
MAX_TOTAL_CODE_CHARS = 200_000    # limit prompt size


def parse_requirements(uploaded_file) -> List[Dict[str, Any]]:
    """
    Parse requirements from an uploaded .docx file.
    Attempts to extract IDs like R-1 or lines starting with numbers/bullets.
    Returns list of dicts: [{"id": "R-1", "text": "..."} , ...]
    """
    try:
        doc = Document(uploaded_file)
    except Exception:
        # If uploaded_file is bytes-like, wrap in BytesIO
        if isinstance(uploaded_file, (bytes, bytearray)):
            doc = Document(io.BytesIO(uploaded_file))
        else:
            raise

    requirements = []
    temp_lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            temp_lines.append(text)

    # Identify requirement lines: those containing pattern R-<number> or starting with a number or bullet
    req_pattern = re.compile(r"^(R-\d+|[A-Za-z]?\d+\.?|\\*|-|•)\\s*(.+)$")
    for line in temp_lines:
        m = req_pattern.match(line)
        if m:
            id_candidate = m.group(1)
            text = m.group(2).strip()
            if not id_candidate.startswith("R-"):
                # Normalize ID
                id_candidate = f"R-{len(requirements)+1}"
            requirements.append({"id": id_candidate, "text": text})

    # Fallback: if no matches, use all lines as sequential requirements
    if not requirements:
        for i, line in enumerate(temp_lines, start=1):
            requirements.append({"id": f"R-{i}", "text": line})

    logger.info("Parsed %d requirements from document", len(requirements))
    return requirements


def read_source_code(directory: str) -> List[Dict[str, Any]]:
    """
    Recursively read source files from the provided directory.
    Returns list of dicts: [{"path": "...", "content": "..."} , ...]
    """
    source_files = []
    total_chars = 0
    for root, _, files in os.walk(directory):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in SUPPORTED_EXTENSIONS:
                path = os.path.join(root, file)
                try:
                    size = os.path.getsize(path)
                    if size > MAX_FILE_SIZE_BYTES:
                        logger.info("Skipping large file: %s (%.2f KB)", path, size/1024)
                        continue
                    with open(path, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                        total_chars += len(content)
                        if total_chars > MAX_TOTAL_CODE_CHARS:
                            logger.info("Reached max total code chars limit. Truncating further additions.")
                            break
                        source_files.append({"path": path, "content": content})
                except Exception as e:
                    logger.exception("Failed reading file %s: %s", path, e)
        else:
            # continue the outer walk even if inner breaks due to char limit
            continue
        # break outer loop if char limit exceeded
        if total_chars > MAX_TOTAL_CODE_CHARS:
            break

    logger.info("Collected %d source files", len(source_files))
    return source_files


def build_structured_prompt(requirements: List[Dict[str, Any]], source_files: List[Dict[str, Any]]) -> str:
    """
    Create a structured prompt for the selected LLM.
    Includes clear instructions and a compact code listing.
    """
    req_text = "\n".join([f"- {r['id']}: {r['text']}" for r in requirements])

    # Compact code listing with file separators
    code_sections = []
    for sf in source_files:
        header = f"FILE: {sf['path']}\n" + ("-" * 60)
        code_sections.append(f"{header}\n{sf['content']}")
    code_text = "\n\n".join(code_sections)

    prompt = f"""
